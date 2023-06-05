# -*- coding: utf-8 -*-
"""
Writes MS Word document summarizing key values from test.

Created on Thu Mar  2 21:28:35 2023

@author: wfros
"""








# %% Imports and Definitions
import Model_10_Story
import docx
from time import localtime, strftime
import os
import json


'''
paragraph = report.add_paragraph('')

report.add_heading('', level=1)
# Title is level 0.

report.add_page_break()

report.add_picture(out_dir + '/base_shear_total.png', width=docx.shared.Inches(6.0))






'''
# %% Main Function
def write_report(analysisID, case):

    # %% Start
    out_dir = Model_10_Story.__path__[0]
    out_dir = out_dir + '/out/' + analysisID + '/' + case
    
    report = docx.Document()
    
    # %%% Style Definitions
    styles = report.styles
    
    style = styles.add_style('TCenter', docx.enum.style.WD_STYLE_TYPE.PARAGRAPH)
    styles['TCenter'].paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    styles['TCenter'].font.size = docx.shared.Pt(24)
    styles['TCenter'].font.underline = True
    
    styles['Title'].font.color.rgb = docx.shared.RGBColor(0, 0, 0)
    styles['Title'].paragraph_format.space_before = docx.shared.Pt(0)
    
    styles['Heading 1'].paragraph_format.space_before = docx.shared.Pt(6)
    
    styles['Heading 2'].paragraph_format.space_before = docx.shared.Pt(0)
    
    styles['Normal'].paragraph_format.space_after = docx.shared.Pt(0)
    
    styles.add_style('Center', docx.enum.style.WD_STYLE_TYPE.PARAGRAPH)
    styles['Center'].paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    styles['Center'].paragraph_format.space_after = docx.shared.Pt(0)
    
    def _center_cells(table):
        for col in table.columns:
            for cell in col.cells:
                cell.paragraphs[0].style = 'Center'
    
    # %%% Import JSON Data
    # Case
    try:
        with open(out_dir + '/outdata.json', 'r') as file:
            data = json.load(file)
    except FileNotFoundError:
        raise FileNotFoundError('JSON file not yet created.')
    
    # Model
    modeldata = Model_10_Story.__path__[0] + '/out/' + analysisID + '/data.json'
    try:
        with open(modeldata, 'r') as file:
            modeldata = json.load(file)
    except FileNotFoundError:
        raise FileNotFoundError('JSON file not yet created.')
    
    
    # %% Write Document
    
    # %%% Header
    report.add_paragraph('          ' + case + '         \u00A0', style='TCenter')
    report.add_heading('Time History Analysis', level=1)
    report.add_paragraph('Date:\t\t' + strftime('%A, %b-%d-%Y %H:%M', localtime()))
    name = os.getlogin()
    name = 'W Roser' if name == 'wfros' else name
    report.add_paragraph('Run by:\t\t' + name)
    
    # %%% Summary of Maximum Values
    report.add_heading('Summary of Maximum Responses', level=1)
    
    # %%%% Drifts along wall Lines
    report.add_heading('Drifts Along Each Wall', level=2)
    report.add_picture(out_dir + '/wall_drift_profile.png', width=docx.shared.Inches(6.0))
    table = report.add_table(rows=12, cols=5, style='Table Grid')
    
    # Header Row
    row = table.rows[0].cells
    row[0].text = 'Level'
    row[0].merge(table.cell(1, 0))
    row[1].text = 'X'
    row[1].merge(row[2])
    row[3].text = 'Y'
    row[3].merge(row[4])
    
    # Secondary Header Row
    row = table.rows[1].cells
    
    
    for i, floor in enumerate(range(11, 1, -1)):
        # Row label
        cell = table.cell(i + 2, 0)
        cell.text = cell.text = 'R' if floor == 11 else str(floor)
        
        for col, wall in enumerate(['North CLT', 'South CLT', 'West MPP', 'East MPP']):
            row[col + 1].text = wall.split(' ')[0]
            # Displacement
            table.cell(i + 2, col + 1).text = '{:.2f}%'.format(data['Max Drift'][wall][str(floor)] * 100)
    
    _center_cells(table)
        
    # %%%% Building Center of Mass
    report.add_heading('Building Center of Mass', level=2)
    table = report.add_table(rows=13, cols=5, style='Table Grid')
    
    # Header Row
    row = table.rows[0].cells
    row[0].text = 'Level'
    row[0].merge(table.cell(1, 0))
    row[1].text = 'Displacement (in.)'
    row[1].merge(row[2])
    row[3].text = 'Acceleration (in./sec²)'
    row[3].merge(row[4])
    
    # Secondary Header Row
    row = table.rows[1].cells
    for col, heading in enumerate('XYXY'):
        row[col + 1].text = heading
    
    for i, floor in enumerate(range(11, 0, -1)):
        # Row label
        cell = table.cell(i + 2, 0)
        cell.text = 'R' if floor == 11 else 'G' if floor == 1 else str(floor)
        
        if not floor == 1:
            # Displacement
            table.cell(i + 2, 1).text = '{:.2f}'.format(data['disp_']['X'][str(floor)])
            table.cell(i + 2, 2).text = '{:.2f}'.format(data['disp_']['Y'][str(floor)])
            # Acceleration
            table.cell(i + 2, 3).text = '{:.0f}'.format(data['accel_']['X'][str(floor)])
            table.cell(i + 2, 4).text = '{:.0f}'.format(data['accel_']['Y'][str(floor)])
            
        else:
            try:
                table.cell(i + 2, 3).text = '{:.0f}'.format(data['PGA']['X'] * 386.4)
                table.cell(i + 2, 4).text = '{:.0f}'.format(data['PGA']['Y'] * 386.4)
            except:
                pass
    
    _center_cells(table)
    
    report.add_paragraph('')
    
    # %%%% Base of Walls
    report.add_heading('Wall Base', level=2)
    table = report.add_table(rows=6, cols=9, style='Table Grid')
    
    # Header Rows
    row = table.rows[0].cells
    row[0].text = 'Wall'
    row[0].merge(table.cell(1, 0))
    headings = [['Shear (kip)', 'X', 'Y'],
                ['Moment (kip-in.)', 'about Y', 'about X'],
                ['Rotation (rad.)', 'about Y', 'about X'],
                ['Toe Uplift (in.)', 'SW', 'NE']]
    for i, heading in enumerate(headings):
        table.cell(0, 1 + 2*i).text = heading[0]
        table.cell(0, 1 + 2*i).merge(table.cell(0, 2 + 2*i))
        table.cell(1, 1 + 2*i).text = heading[1]
        table.cell(1, 2 + 2*i).text = heading[2]
    
    for row, wall in zip(range(2, 6), ['CLT North', 'CLT South', 'MPP West', 'MPP East']):
        table.cell(row, 0).text = wall.split(' ')[-1]
        # Shear
        table.cell(row, 1).text = '{:.1f}'.format(data['Base Shear']['X'][wall])
        table.cell(row, 2).text = '{:.1f}'.format(data['Base Shear']['Y'][wall])
        # Moment
        table.cell(row, 3).text = '{:.0f}'.format(data['Base MR']['Moment']['East-West'][wall[:5]])
        table.cell(row, 4).text = '{:.0f}'.format(data['Base MR']['Moment']['North-South'][wall[:5]])
        # Rotation
        table.cell(row, 5).text = '{:.4f}'.format(data['Base MR']['Rotation']['East-West'][wall[:5]])
        table.cell(row, 6).text = '{:.4f}'.format(data['Base MR']['Rotation']['North-South'][wall[:5]])
        # Toe Uplift
        dir1 = wall.split(' ')[-1]
        dir2 = 'West' if dir1.endswith('th') else 'South'
        table.cell(row, 7).text = '{:.2f}'.format(data['Toe Uplift'][wall[:3]][dir2][dir1])
        dir2 = 'East' if dir1.endswith('th') else 'North'
        table.cell(row, 8).text = '{:.2f}'.format(data['Toe Uplift'][wall[:3]][dir2][dir1])
    
    _center_cells(table)
    
    bs_text = ''
    for direction in 'XY':
        bs_text = bs_text + direction +': {:.1f} kip\t'.format(data['Base Shear'][direction]['Total'])
    report.add_paragraph('Peak Base Shear:\t' + bs_text)
    
    report.add_paragraph('')
    
    # %%%% Eigen
    report.add_heading('Eigen', level=2)
    eigen = modeldata['Eigen']
    table = report.add_table(rows=len(eigen['T']) + 1, cols=len(eigen) + 1, style='Table Grid')
    
    # Header Rows
    headers = ['Mode'] + [x for x in eigen.keys()]
    for row in range(len(eigen['T']) + 1):
        for col in range(len(eigen) + 1):
            if row == 0:
                table.cell(row, col).text = headers[col]
            else:
                if col == 0:
                    table.cell(row, col).text = str(row)
                else:
                    table.cell(row, col).text = '{:.2e}'.format(eigen[headers[col]][row - 1])
    
    _center_cells(table)
    
    report.add_page_break()
    
    # %%% EQ Info
    # Get gm files
    report.add_heading('Earthquake Information', level=1)
    report.add_paragraph('Name:\t' + case)
    
    pga_text = ''
    for direction, pga in data['PGA'].items():
        pga_text = pga_text + direction +': {:.2e} g\t'.format(pga)
    report.add_paragraph('Peak Ground Accelerations:\t' + pga_text)
    report.add_picture(out_dir + '/gm_acceleration.png', width=docx.shared.Inches(6.0))
    
    report.add_heading('Response Spectra', level=2)
    for direction in data['PGA'].keys():
        try:
            report.add_picture(out_dir + '/gm_spectrum_' + direction + '.png', width=docx.shared.Inches(6.0))
        except FileNotFoundError:
            pass
    
    # %%% Individual Charts
    # Drift
    report.add_heading('Responses of Interest', level=1)
    report.add_heading('Drift Histories', level=2)
    report.add_picture(out_dir + '/drift_profile_North.png', width=docx.shared.Inches(6.0))
    report.add_picture(out_dir + '/drift_profile_South.png', width=docx.shared.Inches(6.0))
    report.add_picture(out_dir + '/drift_profile_West.png', width=docx.shared.Inches(6.0))
    report.add_picture(out_dir + '/drift_profile_East.png', width=docx.shared.Inches(6.0))
    
    report.add_page_break()
    
    # Displacement
    report.add_heading('Drift Histories', level=2)
    report.add_picture(out_dir + '/disp_upper.png', width=docx.shared.Inches(6.0))
    report.add_picture(out_dir + '/disp_lower.png', width=docx.shared.Inches(6.0))
    
    report.add_page_break()
    
    # Acceleration
    report.add_heading('Acceleration Histories', level=2)
    report.add_picture(out_dir + '/accel_upper.png', width=docx.shared.Inches(6.0))
    report.add_picture(out_dir + '/accel_lower.png', width=docx.shared.Inches(6.0))
    
    report.add_page_break()
    
    # Base Shear, Moment, Uplift
    report.add_heading('Base Shear', level=2)
    report.add_heading('Entire Building', level=3)
    report.add_picture(out_dir + '/base_shear_total.png', width=docx.shared.Inches(6.0))
    report.add_heading('Base of Each Wall', level=3)
    report.add_picture(out_dir + '/base_shear_walls.png', width=docx.shared.Inches(6.0))
    
    report.add_page_break()
    
    report.add_heading('Base Moment-Rotation', level=2)
    report.add_paragraph('Sign of values has been changed from the model for aesthetics.')
    report.add_picture(out_dir + '/base_moment_East-West.png', width=docx.shared.Inches(6.0))
    report.add_picture(out_dir + '/base_moment_North-South.png', width=docx.shared.Inches(6.0))
    
    report.add_page_break()
    
    report.add_heading('Toe Uplift', level=2)
    report.add_paragraph('Each subfigure represents the same toe on parallel walls. Overlapping lines indicate low torsion.')
    report.add_picture(out_dir + '/toe_uplift_CLT.png', width=docx.shared.Inches(6.0))
    report.add_picture(out_dir + '/toe_uplift_MPP.png', width=docx.shared.Inches(6.0))
    
    report.add_page_break()
    
    report.add_heading('Base Strains', level=2)
    report.add_paragraph('Each subfigure represents the same toe on parallel walls. Overlapping lines indicate low torsion.')
    report.add_picture(out_dir + '/toe_uplift_CLT.png', width=docx.shared.Inches(6.0))
    report.add_picture(out_dir + '/toe_uplift_MPP.png', width=docx.shared.Inches(6.0))
    
    report.add_page_break()
    
    # PT
    report.add_heading('PT Axial Force', level=2)
    report.add_picture(out_dir + '/PT_axial_CLT North.png', width=docx.shared.Inches(6.0))
    report.add_picture(out_dir + '/PT_axial_CLT North.png', width=docx.shared.Inches(6.0))
    report.add_picture(out_dir + '/PT_axial_CLT North.png', width=docx.shared.Inches(6.0))
    report.add_picture(out_dir + '/PT_axial_CLT North.png', width=docx.shared.Inches(6.0))
    
    report.add_heading('PT Lateral Force', level=2)
    report.add_picture(out_dir + '/PT_total.png', width=docx.shared.Inches(6.0))
    report.add_picture(out_dir + '/PT_each_wall.png', width=docx.shared.Inches(6.0))
    
    
    # %% Save
    report.save(out_dir + '/' + case + '.docx')


# %% Run as Script
if __name__ == '__main__':
    # Output Folder Selection
    analysisID = 'Report_Test'
    case = '27_Northridge_MCE'
    
    
    write_report(analysisID, case)